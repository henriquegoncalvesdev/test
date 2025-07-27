from flask import Flask, request, jsonify, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_cors import CORS
from flask_migrate import Migrate
from datetime import datetime
import os
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches
import json
import re

app = Flask(__name__)
app.config['SECRET_KEY'] = 'sua_chave_secreta_aqui'
app.config['SQLALCHEMY_DATABASE_URI'] = 'mysql+pymysql://root@localhost/cliente_system'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Criar pasta de uploads se não existir
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

db = SQLAlchemy(app)
migrate = Migrate(app, db)
CORS(app)

# Modelos do banco de dados
class Cliente(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    nome_completo = db.Column(db.String(200), nullable=False)
    cpf_cnpj = db.Column(db.String(20), unique=True, nullable=False)
    endereco = db.Column(db.Text)
    telefone = db.Column(db.String(20))
    email = db.Column(db.String(100))
    data_cadastro = db.Column(db.DateTime, default=datetime.utcnow)
    observacoes = db.Column(db.Text)
    
    # Relacionamentos
    documentos = db.relationship('Documento', backref='cliente', lazy=True)
    custeios = db.relationship('Custeio', backref='cliente', lazy=True)
    investimentos = db.relationship('Investimento', backref='cliente', lazy=True)

class Documento(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    cliente_id = db.Column(db.Integer, db.ForeignKey('cliente.id'), nullable=False)
    nome_arquivo = db.Column(db.String(255), nullable=False)
    tipo_documento = db.Column(db.String(50))  # 'custeio', 'investimento', 'outro'
    data_upload = db.Column(db.DateTime, default=datetime.utcnow)
    caminho_arquivo = db.Column(db.String(500))
    observacoes = db.Column(db.Text)

class Custeio(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    cliente_id = db.Column(db.Integer, db.ForeignKey('cliente.id'), nullable=False)
    valor = db.Column(db.Numeric(15, 2))
    finalidade = db.Column(db.Text)
    prazo = db.Column(db.String(50))
    taxa_juros = db.Column(db.String(20))
    garantia = db.Column(db.Text)
    data_solicitacao = db.Column(db.DateTime, default=datetime.utcnow)
    status = db.Column(db.String(20), default='pendente')  # pendente, aprovado, rejeitado
    observacoes = db.Column(db.Text)

class Investimento(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    cliente_id = db.Column(db.Integer, db.ForeignKey('cliente.id'), nullable=False)
    valor_investimento = db.Column(db.Numeric(15, 2))
    tipo_investimento = db.Column(db.String(100))
    prazo_investimento = db.Column(db.String(50))
    rentabilidade = db.Column(db.String(20))
    objetivo = db.Column(db.Text)
    data_inicio = db.Column(db.DateTime, default=datetime.utcnow)
    status = db.Column(db.String(20), default='ativo')  # ativo, finalizado, cancelado
    observacoes = db.Column(db.Text)

# Rotas da API
@app.route('/api/clientes', methods=['GET'])
def listar_clientes():
    clientes = Cliente.query.all()
    return jsonify([{
        'id': c.id,
        'nome_completo': c.nome_completo,
        'cpf_cnpj': c.cpf_cnpj,
        'endereco': c.endereco,
        'telefone': c.telefone,
        'email': c.email,
        'data_cadastro': c.data_cadastro.isoformat() if c.data_cadastro else None,
        'observacoes': c.observacoes
    } for c in clientes])

@app.route('/api/clientes', methods=['POST'])
def criar_cliente():
    data = request.json
    novo_cliente = Cliente(
        nome_completo=data['nome_completo'],
        cpf_cnpj=data['cpf_cnpj'],
        endereco=data.get('endereco'),
        telefone=data.get('telefone'),
        email=data.get('email'),
        observacoes=data.get('observacoes')
    )
    db.session.add(novo_cliente)
    db.session.commit()
    return jsonify({'id': novo_cliente.id, 'message': 'Cliente criado com sucesso'}), 201

@app.route('/api/clientes/<int:cliente_id>', methods=['GET'])
def obter_cliente(cliente_id):
    cliente = Cliente.query.get_or_404(cliente_id)
    return jsonify({
        'id': cliente.id,
        'nome_completo': cliente.nome_completo,
        'cpf_cnpj': cliente.cpf_cnpj,
        'endereco': cliente.endereco,
        'telefone': cliente.telefone,
        'email': cliente.email,
        'data_cadastro': cliente.data_cadastro.isoformat() if cliente.data_cadastro else None,
        'observacoes': cliente.observacoes
    })

@app.route('/api/clientes/<int:cliente_id>', methods=['PUT'])
def atualizar_cliente(cliente_id):
    cliente = Cliente.query.get_or_404(cliente_id)
    data = request.json
    cliente.nome_completo = data['nome_completo']
    cliente.cpf_cnpj = data['cpf_cnpj']
    cliente.endereco = data.get('endereco')
    cliente.telefone = data.get('telefone')
    cliente.email = data.get('email')
    cliente.observacoes = data.get('observacoes')
    db.session.commit()
    return jsonify({'message': 'Cliente atualizado com sucesso'})

@app.route('/api/clientes/<int:cliente_id>', methods=['DELETE'])
def deletar_cliente(cliente_id):
    cliente = Cliente.query.get_or_404(cliente_id)
    db.session.delete(cliente)
    db.session.commit()
    return jsonify({'message': 'Cliente deletado com sucesso'})

# Rotas para documentos
@app.route('/api/clientes/<int:cliente_id>/documentos', methods=['GET'])
def listar_documentos(cliente_id):
    documentos = Documento.query.filter_by(cliente_id=cliente_id).all()
    return jsonify([{
        'id': d.id,
        'nome_arquivo': d.nome_arquivo,
        'tipo_documento': d.tipo_documento,
        'data_upload': d.data_upload.isoformat() if d.data_upload else None,
        'observacoes': d.observacoes
    } for d in documentos])

@app.route('/api/clientes/<int:cliente_id>/documentos', methods=['POST'])
def upload_documento(cliente_id):
    if 'arquivo' not in request.files:
        return jsonify({'error': 'Nenhum arquivo enviado'}), 400
    
    arquivo = request.files['arquivo']
    if arquivo.filename == '':
        return jsonify({'error': 'Nenhum arquivo selecionado'}), 400
    
    filename = secure_filename(arquivo.filename)
    caminho_arquivo = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    arquivo.save(caminho_arquivo)
    
    novo_documento = Documento(
        cliente_id=cliente_id,
        nome_arquivo=filename,
        tipo_documento=request.form.get('tipo_documento', 'outro'),
        caminho_arquivo=caminho_arquivo,
        observacoes=request.form.get('observacoes')
    )
    db.session.add(novo_documento)
    db.session.commit()
    
    return jsonify({'message': 'Documento enviado com sucesso'})

# Rotas para custeio
@app.route('/api/clientes/<int:cliente_id>/custeios', methods=['GET'])
def listar_custeios(cliente_id):
    custeios = Custeio.query.filter_by(cliente_id=cliente_id).all()
    return jsonify([{
        'id': c.id,
        'valor': float(c.valor) if c.valor else None,
        'finalidade': c.finalidade,
        'prazo': c.prazo,
        'taxa_juros': c.taxa_juros,
        'garantia': c.garantia,
        'data_solicitacao': c.data_solicitacao.isoformat() if c.data_solicitacao else None,
        'status': c.status,
        'observacoes': c.observacoes
    } for c in custeios])

@app.route('/api/clientes/<int:cliente_id>/custeios', methods=['POST'])
def criar_custeio(cliente_id):
    data = request.json
    novo_custeio = Custeio(
        cliente_id=cliente_id,
        valor=data.get('valor'),
        finalidade=data.get('finalidade'),
        prazo=data.get('prazo'),
        taxa_juros=data.get('taxa_juros'),
        garantia=data.get('garantia'),
        observacoes=data.get('observacoes')
    )
    db.session.add(novo_custeio)
    db.session.commit()
    return jsonify({'id': novo_custeio.id, 'message': 'Custeio criado com sucesso'}), 201

# Rotas para investimento
@app.route('/api/clientes/<int:cliente_id>/investimentos', methods=['GET'])
def listar_investimentos(cliente_id):
    investimentos = Investimento.query.filter_by(cliente_id=cliente_id).all()
    return jsonify([{
        'id': i.id,
        'valor_investimento': float(i.valor_investimento) if i.valor_investimento else None,
        'tipo_investimento': i.tipo_investimento,
        'prazo_investimento': i.prazo_investimento,
        'rentabilidade': i.rentabilidade,
        'objetivo': i.objetivo,
        'data_inicio': i.data_inicio.isoformat() if i.data_inicio else None,
        'status': i.status,
        'observacoes': i.observacoes
    } for i in investimentos])

@app.route('/api/clientes/<int:cliente_id>/investimentos', methods=['POST'])
def criar_investimento(cliente_id):
    data = request.json
    novo_investimento = Investimento(
        cliente_id=cliente_id,
        valor_investimento=data.get('valor_investimento'),
        tipo_investimento=data.get('tipo_investimento'),
        prazo_investimento=data.get('prazo_investimento'),
        rentabilidade=data.get('rentabilidade'),
        objetivo=data.get('objetivo'),
        observacoes=data.get('observacoes')
    )
    db.session.add(novo_investimento)
    db.session.commit()
    return jsonify({'id': novo_investimento.id, 'message': 'Investimento criado com sucesso'}), 201

# Rota para gerar documentos Word
@app.route('/api/gerar-documento', methods=['POST'])
def gerar_documento():
    data = request.json
    tipo_documento = data['tipo']  # 'custeio' ou 'investimento'
    cliente_id = data['cliente_id']
    
    cliente = Cliente.query.get_or_404(cliente_id)
    
    # Criar documento Word
    doc = Document()
    
    if tipo_documento == 'custeio':
        # Buscar dados do custeio se existir
        custeio = Custeio.query.filter_by(cliente_id=cliente_id).first()
        
        # Cabeçalho
        doc.add_heading('DOCUMENTO DE CUSTEIO', 0)
        doc.add_paragraph()
        
        # Informações do cliente
        doc.add_heading('INFORMAÇÕES DO CLIENTE', level=1)
        doc.add_paragraph(f'Nome Completo: {cliente.nome_completo}')
        doc.add_paragraph(f'CPF/CNPJ: {cliente.cpf_cnpj}')
        doc.add_paragraph(f'Endereço: {cliente.endereco or "Não informado"}')
        doc.add_paragraph(f'Telefone: {cliente.telefone or "Não informado"}')
        doc.add_paragraph(f'Email: {cliente.email or "Não informado"}')
        doc.add_paragraph()
        
        # Informações do custeio
        if custeio:
            doc.add_heading('INFORMAÇÕES DO CUSTEIO', level=1)
            doc.add_paragraph(f'Valor: R$ {custeio.valor:,.2f}' if custeio.valor else 'Valor: Não informado')
            doc.add_paragraph(f'Finalidade: {custeio.finalidade or "Não informado"}')
            doc.add_paragraph(f'Prazo: {custeio.prazo or "Não informado"}')
            doc.add_paragraph(f'Taxa de Juros: {custeio.taxa_juros or "Não informado"}')
            doc.add_paragraph(f'Garantia: {custeio.garantia or "Não informado"}')
            doc.add_paragraph()
        
        # Data de geração
        doc.add_paragraph(f'Documento gerado em: {datetime.now().strftime("%d/%m/%Y às %H:%M")}')
        
    elif tipo_documento == 'investimento':
        # Buscar dados do investimento se existir
        investimento = Investimento.query.filter_by(cliente_id=cliente_id).first()
        
        # Cabeçalho
        doc.add_heading('DOCUMENTO DE INVESTIMENTO', 0)
        doc.add_paragraph()
        
        # Informações do cliente
        doc.add_heading('INFORMAÇÕES DO CLIENTE', level=1)
        doc.add_paragraph(f'Nome Completo: {cliente.nome_completo}')
        doc.add_paragraph(f'CPF/CNPJ: {cliente.cpf_cnpj}')
        doc.add_paragraph(f'Endereço: {cliente.endereco or "Não informado"}')
        doc.add_paragraph(f'Telefone: {cliente.telefone or "Não informado"}')
        doc.add_paragraph(f'Email: {cliente.email or "Não informado"}')
        doc.add_paragraph()
        
        # Informações do investimento
        if investimento:
            doc.add_heading('INFORMAÇÕES DO INVESTIMENTO', level=1)
            doc.add_paragraph(f'Valor do Investimento: R$ {investimento.valor_investimento:,.2f}' if investimento.valor_investimento else 'Valor: Não informado')
            doc.add_paragraph(f'Tipo de Investimento: {investimento.tipo_investimento or "Não informado"}')
            doc.add_paragraph(f'Prazo do Investimento: {investimento.prazo_investimento or "Não informado"}')
            doc.add_paragraph(f'Rentabilidade: {investimento.rentabilidade or "Não informado"}')
            doc.add_paragraph(f'Objetivo: {investimento.objetivo or "Não informado"}')
            doc.add_paragraph()
        
        # Data de geração
        doc.add_paragraph(f'Documento gerado em: {datetime.now().strftime("%d/%m/%Y às %H:%M")}')
    
    # Salvar documento
    filename = f"{tipo_documento}_{cliente.nome_completo.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    caminho_arquivo = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    doc.save(caminho_arquivo)
    
    return send_file(caminho_arquivo, as_attachment=True, download_name=filename)

# Rota para estatísticas
@app.route('/api/estatisticas', methods=['GET'])
def obter_estatisticas():
    total_clientes = Cliente.query.count()
    total_custeios = Custeio.query.count()
    total_investimentos = Investimento.query.count()
    total_documentos = Documento.query.count()
    
    return jsonify({
        'total_clientes': total_clientes,
        'total_custeios': total_custeios,
        'total_investimentos': total_investimentos,
        'total_documentos': total_documentos
    })

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run(debug=True, host='0.0.0.0', port=5001) 